from inspect import getmembers, isclass, isabstract
from . import generators as generators_module
from .. import doctools
import pandas as pd
import docx
import os.path


def _delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def _delete_table(table):
    table._element.getparent().remove(table._element)


class Orchestrator:
    """
    Orchestration object to run documentation generation.
    Uses data and metadata passed during init with generators
    to produce outputs to a .docx word file.
    """

    STYLE_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates_docx", "style_template_365_empty.docx")
    
    def __init__(self, file, data=None, metadata=None, output_folder=None):
        """Creates a orchestrator object with relevant data and metadata.

        Args:
            file (str): filename of .docx for output (include .docx in name)
            output_folder (str): the output folder for the documentation files
            data (pd.DataFrame, optional): Defaults to None.
            metadata (dict, optional): Dictionary containing metadata - such as 
            columns names for target, predictors or filepaths for models. Defaults to None.
        """
        self.data = data
        self.metadata = metadata
        self.file = file
        self.output_folder = output_folder

        self._transform_predictors()
        self._load_generators()
        self._init_doctools()

    def _load_generators(self):
        self.generators = {}
        for _, member in getmembers(generators_module):
            if isclass(member):
                if not isabstract(member) and issubclass(member, generators_module.Generator):
                    self.generators[member.ID] = member

    def _print_generators(self):
        for ID, gen in self.generators.items():
            print(ID, "\n", gen.__doc__)

    def _transform_predictors(self):
        if 'predictors' in self.metadata.keys():
            for pred in self.metadata['predictors']:
                if not pd.api.types.is_numeric_dtype(self.data[pred].dtype):
                    self.data[pred] = self.data[pred].astype('category')

    def _init_doctools(self):
        self.doctools = doctools.ProjectParameters()

        columns = self.metadata["columns"]
        self.doctools.targets = [(columns["target"], columns["base"]), (columns["short_target"], columns["short_base"])]
        self.doctools.time_variable = columns["time"]
        self.doctools.rowid_variable = columns["row_id"]
        if "use_weight" in self.metadata.keys():
            self.use_weight = self.metadata["use_weight"]
            if self.use_weight:
                self.doctools.weight = columns["weight"]
        else:
            self.use_weight = False

        self.doctools.sample_dict = self.metadata["samples"]

    def _init_docx(self):
        self.document = docx.Document(self.STYLE_TEMPLATE_PATH)
        for paragraph in self.document.paragraphs:
            _delete_paragraph(paragraph)

        for table in self.document.tables:
            _delete_table(table)

    def _add_to_structure(self, line, predictor_name=None):
        processed_line = line.lstrip("-").split(":")
        if len(processed_line) > 1:
            heading = processed_line[1]
        else:
            heading = None
        self.structure.append({
            "generator": processed_line[0],
            "level": len(line) - len(line.lstrip("-")),
            "heading": heading,
            "predictor_name": predictor_name,
        })

        if processed_line[0] not in self.generators.keys():
            print(f"{processed_line[0]} is not a valid generator ID")

    def _process_predictor_line_cache(self, cache, predictor_names):
        for predictor_name in predictor_names:
            for line in cache:
                self._add_to_structure(line, predictor_name)

    def _validate_structure(self):
        self.structure = []
        predictor_loop = False
        predictor_line_cache = []
        
        for line in self._raw_structure:
            if line[0] == '@' and not predictor_loop:
                predictor_loop = True
                loop_defining_list_name = line[1:]
            elif line[0] == '@' and predictor_loop:
                self._process_predictor_line_cache(predictor_line_cache, self.metadata[loop_defining_list_name])
                predictor_loop = False
                predictor_line_cache = []
            elif not predictor_loop:
                self._add_to_structure(line)
            else:
                predictor_line_cache.append(line)

    def load_structure(self, file_path):
        """Loads a structure for documentation from file_path.
        Structure is then validated.

        Args:
            file_path (str): relative or absolute path to file
        """
        with open(file_path, "r") as file:
            self._raw_structure = [line.strip() for line in file.readlines() if line[0] != "#"]

        self._validate_structure()

    def _create_generator_instances(self):
        self._document_elements = []
        for structure_dict in self.structure:
            g = self.generators[structure_dict["generator"]](
                self,
                structure_dict["level"],
                heading=structure_dict["heading"],
                predictor_name=structure_dict["predictor_name"],
                output_folder=self.output_folder
            )
            g.calculate()

            self._document_elements.append(g)

    def _export_all_to_doc(self):
        for g in self._document_elements:
            g.export_to_document()

    def _save_document(self):
        file_final = self.file
        num_suffix = 0
        while os.path.exists(f"{file_final}.docx"):
            num_suffix += 1
            file_final = self.file + f"_{str(num_suffix)}"
        self.document.save(f"{file_final}.docx")

    def export(self):
        """Initialized the docx document and runs through all 
        generator to produce outputs and export them to spefied
        file.
        """
        self._init_docx()
        self._create_generator_instances()
        self._export_all_to_doc()
        self._save_document()