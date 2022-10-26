from abc import ABC, abstractmethod
import numpy as np
import pandas as pd
import pickle
import os
import json
from .. import grouping_new
from .. import grouping
from .. import scorecard
from .. import features
from .. import metrics
from datetime import date
import lightgbm as lgb
from sklearn.linear_model import LogisticRegression

from os import path
import docx


class Generator(ABC):
    """Template class for documentation generators.

    All generators should inherit from this base class,
    define a unique class.ID and implement following methods:

    __init__(...):
        extend the base init and load any needed information
        from orchestrator.metadata

    calculate():
        using data and metadata from orchestrator, calculate
        all required output and/or generate plots and other
        outputs

    export_to_document():
        save outputs to .docx file object stored
        in orchestrator object

    """
    DEFAULT_HEADING = ""
    PIC_WIDTH_S = 3.25
    PIC_WIDTH_M = 5.0
    PIC_WIDTH_L = 6.5

    def __init__(self, orchestrator, level, heading=None, predictor_name=None, output_folder=None):
        self.orchestrator = orchestrator
        self.level = level
        self.columns = orchestrator.metadata["columns"]
        self.predictors = orchestrator.metadata["predictors"]
        self.data = self.orchestrator.data

        if heading:
            self.heading = heading
        else:
            self.heading = self.DEFAULT_HEADING
        if "use_weight" in self.orchestrator.metadata.keys():
            self.use_weight = self.orchestrator.metadata["use_weight"]
        else:
            self.use_weight = False
        if predictor_name:
            self.predictor_name = predictor_name
        else:
            self.predictor_name = None
        if output_folder:
            self.output_folder = output_folder
        else:
            self.output_folder = ""

    
    def table_from_dataframe(self, dataframe, style=None, float_precision=None, print_index=True):
        ## TO DO
        # multi index can cause ugly formatting (ref data summary)

        # add a table to the end and create a reference variable
        # extra row and column are for index and column names
        if not print_index:
            index_width = 0
        elif isinstance(dataframe.index, pd.MultiIndex):
            index_width = len(dataframe.index.levels)
        else:
            index_width = 1

        columns_index_height = 1

        table = self.orchestrator.document.add_table(dataframe.shape[0] + columns_index_height, dataframe.shape[1] + index_width)
        table.style = self.orchestrator.document.styles["Table Grid"]

        # add the header rows.
        for column_index, column_name in enumerate(dataframe.columns, start=index_width):
            if type(column_name) is tuple:
                column_name = "\n".join(column_name)
            else:
                column_name = str(column_name)

            table.cell(0, column_index).text = column_name

        # add index values
        if print_index:
            if isinstance(dataframe.index, pd.MultiIndex):
                for i, row_name in enumerate(dataframe.index, start=1):
                    for j in range(0, index_width):
                        row_name = dataframe.index[i - 1][j]
                        table.cell(i, j).text = str(row_name)
                        table.cell(i, j).paragraphs[0].runs[0].font.bold = True

            else:
                for i, row_name in enumerate(dataframe.index, start=1):
                    table.cell(i, 0).text = str(row_name)

        # add the rest of the data frame
        for i in range(dataframe.shape[0]):
            for j in range(dataframe.shape[1]):
                value = dataframe.values[i, j]
                if float_precision:
                    try:
                        value = np.round(value, float_precision)
                    except:
                        pass

                table.cell(i + columns_index_height, j + index_width).text = str(value)

        return table

    @abstractmethod
    def calculate(self):
        pass

    @abstractmethod
    def export_to_document(self):
        pass


class TitleInfo(Generator):
    """Generates a chapter containing information about author.

    Required metadata:
        "author_name": (str)
    
    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """

    ID = "title"
    DEFAULT_HEADING = "Title"

    def calculate(self):
        """ Sets the author_name as attribute. """
        self.author_name = self.orchestrator.metadata["author_name"]

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_paragraph(f"Author: {self.author_name}")


class ScoreCardInfo(Generator):
    """Generates a chapter containing information about scorecard segment.

    Required metadata:
        "scorecard_name": (str)

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "info_about_scorecard"
    DEFAULT_HEADING = "Scorecard information"

    def calculate(self):
        """ Sets the attribute scorecard_name. """
        self.scorecard_name = self.orchestrator.metadata["scorecard_name"]

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_paragraph(f"Segment: {self.scorecard_name}")


class SampleStatistics(Generator):
    """
        Shows statistics of the dataset - counts of contracts by time interval and badrate.
        Args:
            orchestrator (Orchestrator): documentation orchestrator with metadata
            level (int): chapter depth (level in multi-leveled chapter tree)
            heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
    """
    ID = 'sample_statistics'
    DEFAULT_HEADING = "Sample Statistics"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.target = self.columns["target"]
        self.sample = "Observable"
        self.plot_path = "data_badrate.png"

    def calculate(self):
        """ prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.PlotDataset`"""
        self.orchestrator.doctools.PlotDataset(data=self.data,
                                               sample=self.sample,
                                               target=self.target,
                                               segment_col=None,
                                               use_weight=self.use_weight,
                                               output_folder=self.output_folder,
                                               filename=self.plot_path,
                                               show_plot=False)

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"Sample:\t{self.sample}\n")
        p.add_run(f"Target:\t\t{self.target}\n")
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path), width=docx.shared.Inches(self.PIC_WIDTH_S),)


class LibraryVersion(Generator):
    """
        Adds versions of all required libraries' from requirements
        Args:
            orchestrator (Orchestrator): documentation orchestrator with metadata
            level (int): chapter depth (level in multi-leveled chapter tree)
            heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
    """

    ID = 'library_version'
    DEFAULT_HEADING = "Library Version"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.psw_version = self.orchestrator.metadata["PSW_version"]
        self.filename = 'libraries_version.csv'

    def calculate(self):
        """ prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.LibrariesVersion`"""
        self.library_version = self.orchestrator.doctools.LibrariesVersion(psw_version=self.psw_version,
                                                                           output_folder=self.output_folder,
                                                                           filename=self.filename)

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"Python Score Version:\t{self.psw_version}\n")
        self.table_from_dataframe(self.library_version, style=None)


class CalibrationPlot(Generator):
    """Creates calibration plot.

        Args:
            orchestrator (Orchestrator): documentation orchestrator with metadata
            level (int): chapter depth (level in multi-leveled chapter tree)
            heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
    """
    ID = "calibration_plot"
    DEFAULT_HEADING = "Calibration Plot"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.score = self.columns["score"]
        self.target = self.columns["target"]
        self.sample_t = "Train"
        self.sample_v = "Validation"
        self.plot_path = "calibration.png"
        self.plot_path_v = "calibration_v.png"
        if self.use_weight:
            self.weight = self.columns["weight"]
            self.plot_path_w = "calibration_w.png"
            self.plot_path_v_w = "calibration_v_w.png"

    def calculate(self):
        """ prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.ScoreCalibration`"""
        self.orchestrator.doctools.ScoreCalibration(self.data,
                                                    self.score,
                                                    self.sample_t,
                                                    self.target,
                                                    output_folder=self.output_folder,
                                                    filename=self.plot_path,
                                                    use_weight=False,
                                                    show_plot=False
                                                    )
        self.orchestrator.doctools.ScoreCalibration(self.data,
                                                    self.score,
                                                    self.sample_v,
                                                    self.target,
                                                    output_folder=self.output_folder,
                                                    filename=self.plot_path_v,
                                                    use_weight=False,
                                                    show_plot=False
                                                    )
        if self.use_weight:
            self.orchestrator.doctools.ScoreCalibration(self.data,
                                                        self.score,
                                                        self.sample_t,
                                                        self.target,
                                                        output_folder=self.output_folder,
                                                        filename=self.plot_path_w,
                                                        use_weight=self.use_weight,
                                                        show_plot=False
                                                        )
            self.orchestrator.doctools.ScoreCalibration(self.data,
                                                        self.score,
                                                        self.sample_v,
                                                        self.target,
                                                        output_folder=self.output_folder,
                                                        filename=self.plot_path_v_w,
                                                        use_weight=self.use_weight,
                                                        show_plot=False
                                                        )
    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"Score:\t\t{self.score}\n")
        p.add_run(f"Target:\t\t{self.target}\n")
        p.add_run(f"Sample:\t{self.sample_t}\n")
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path), width=docx.shared.Inches(self.PIC_WIDTH_S),)
        p.add_run(f"Sample:\t{self.sample_v}\n")
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_v), width=docx.shared.Inches(self.PIC_WIDTH_S),)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted Calibration Plot (weight: {self.weight})\n")
            p.add_run(f"Sample:\t{self.sample_t}\n")
            self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_w), width=docx.shared.Inches(self.PIC_WIDTH_S),)
            p.add_run(f"Sample:\t{self.sample_v}\n")
            self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_v_w), width=docx.shared.Inches(self.PIC_WIDTH_S),)


class CalibrationPlotOldScore(CalibrationPlot):
    """ Creates calibration plot for old score, if given. Does not have weighted version.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.ScoreCalibration`
        export_to_document(): export the output to given Word docx
    """
    ID = "calibration_plot_oldscore"
    DEFAULT_HEADING = "Calibration Plot Comparison"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.score = self.columns["old_score"]
        self.plot_path = "calibration_oldscore.png"


class DataSummary(Generator):
    """Generates basic data summary.

    Required metadata:
        "samples": dictionary of doctools samples
        "target": (str)
        "data_type": (str) column with sample name for each row

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "data_summary"
    DEFAULT_HEADING = "Data Summary"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.data = self.orchestrator.data[self.orchestrator.metadata["samples"]["All"]]
        self.target = self.columns["target"]
        self.data_type = self.columns["data_type"]

    def calculate(self):
        """prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.DataSampleSummary`
        """
        self.summary_table = self.orchestrator.doctools.DataSampleSummary(self.data, "All", self.target,
                                                                          segment_col=self.data_type,
                                                                          use_weight=False,)
        if self.use_weight:
            self.summary_table_weight = self.orchestrator.doctools.DataSampleSummary(self.data, "All", self.target,
                                                                            segment_col=self.data_type,
                                                                            use_weight=True)

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.table_from_dataframe(self.summary_table, style=None, float_precision=3, print_index=False)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"\nWeighted:")
            self.table_from_dataframe(self.summary_table_weight, style=None, float_precision=3, print_index=False)


class JustHeading(Generator):
    """Generates a heading with no text in subsequent chapter. Suitable for situation when we want to have just
    subchapters under this heading. Text of the heading is specified in arguments.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "heading"
    DEFAULT_HEADING = "Placeholder heading"

    def calculate(self):
        pass

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)


class EmptyChapter(Generator):
    """Creates a empty chapter with placeholder text to be filled in manually in the Word document.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. efaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "empty_chapter"
    DEFAULT_HEADING = "Placeholder chapter"

    def calculate(self):
        self.text = "Please fill in manually."

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"{self.text}")


class ScorecardTable(Generator):
    """Table with scorecard. Rows of the table represent categories (bins) of predictors.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "scorecard_table"
    DEFAULT_HEADING = "Scorecard Table"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.target = self.columns["target"]
        self.grouping_path = self.orchestrator.metadata["grouping_path"]
        self.model_path = self.orchestrator.metadata["model_path"]
        self.data = self.data[self.orchestrator.metadata["samples"]["Train"]]
        if self.use_weight:
            self.weight = self.columns["weight"]
        else:
            self.weight = None

    def calculate(self):
        gr = grouping_new.interactive.NewGrouping()
        gr.init_underlying_grouping(self.data[self.predictors])
        gr.load(self.grouping_path)
        md = pickle.load(open(self.model_path, 'rb'))
        sc = scorecard.ScoreCard(
            grouping=gr.grouping,
            predictors=md.predictors,
            coefficients=md.coef,
            intercept=md.intercept,
        )

        self.sc_table = sc.scorecard_table_full(
            data=self.data,
            mask=pd.Series(True, index=self.data.index),
            target=self.target,
            weightcol=self.weight,
        )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Scorecard evaluation was performed using weighted data.")
        self.table_from_dataframe(self.sc_table, style=None, float_precision=3, print_index=False)


class ScorecardTableGM(Generator):
    """Table with scorecard. Rows of the table represent categories (bins) of predictors.
    This version is for General model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "scorecard_table_gm"
    DEFAULT_HEADING = "Scorecard Table"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.target = self.columns["target"]
        self.model_path = self.orchestrator.metadata["model_path"]

    def calculate(self):
        model = pickle.load(open(self.model_path, 'rb'))
        self.sc_table = model.scorecard_table
        self.imp_table = model.imputation_table

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.table_from_dataframe(self.sc_table, style=None, float_precision=3, print_index=True)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"\nImputation table:")
        self.table_from_dataframe(self.imp_table, style=None, float_precision=3, print_index=False)


class MarginalContributionRemove(Generator):
    """Table with marginal contribution of each predictor of the model.
    Version for Logistic Regression PSW model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "marginal_contribution_remove"
    DEFAULT_HEADING = "Marginal Contribution of Predictors"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.target = self.columns["target"]
        self.grouping_path = self.orchestrator.metadata["grouping_path"]
        self.model_path = self.orchestrator.metadata["model_path"]
        self.data_train = self.data[self.orchestrator.metadata["samples"]["Train"]]
        self.data_validation = self.data[self.orchestrator.metadata["samples"]["Validation"]]

        if self.use_weight:
            self.weight = self.columns["weight"]
        else:
            self.weight = None

    def calculate(self):
        gr = grouping_new.interactive.NewGrouping()
        gr.init_underlying_grouping(self.data_train[self.predictors])
        gr.load(self.grouping_path)
        md = pickle.load(open(self.model_path, 'rb'))

        X = gr.transform(self.data_train[self.predictors])
        X_valid = gr.transform(self.data_validation[self.predictors])

        if self.use_weight:
            sample_weight = self.data_train[self.weight]
            sample_weight_valid = self.data_validation[self.weight]
        else:
            sample_weight, sample_weight_valid = None, None

        if md.use_cv:
            self.text = "Marginal contribution calculated using cross-validation.\n"
        else:
            self.text = "Marginal contribution calculated using validation data sample.\n"

        self.mc_table = md.marginal_contribution(
            X=X[md.predictors],
            y=self.data_train[self.target],
            X_valid=X_valid[md.predictors],
            y_valid=self.data_validation[self.target],
            sample_weight=sample_weight,
            sample_weight_valid=sample_weight_valid,
            silent=True,
        )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"{self.text}")
        if self.use_weight:
            p.add_run(f"Scorecard evaluation was performed using weighted data.\n")
        self.table_from_dataframe(self.mc_table, style=None, float_precision=3, print_index=False)


class MarginalContributionRemoveGM(Generator):
    """Table with marginal contribution of each predictor of the model.
    Version for General model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "marginal_contribution_remove_gm"
    DEFAULT_HEADING = "Marginal Contribution of Predictors"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.target = self.columns["target"]
        self.model_path = self.orchestrator.metadata["model_path"]
        self.data_train = self.data[self.orchestrator.metadata["samples"]["Train"]]
        self.data_validation = self.data[self.orchestrator.metadata["samples"]["Validation"]]

        if self.use_weight:
            self.weight = self.columns["weight"]
        else:
            self.weight = None

    def calculate(self):
        md = pickle.load(open(self.model_path, 'rb'))

        if self.use_weight:
            sample_weight = self.data_train[self.weight]
            sample_weight_valid = self.data_validation[self.weight]
        else:
            sample_weight, sample_weight_valid = None, None

        self.mc_table = md.marginal_contribution(
            X = self.data_train,
            y = self.data_train[self.target],
            X_valid = self.data_validation,
            y_valid = self.data_validation[self.target],
            w = sample_weight,
            w_valid = sample_weight_valid,
        )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        if self.use_weight:
            p.add_run(f"Scorecard evaluation was performed using weighted data.\n")
        self.table_from_dataframe(self.mc_table, style=None, float_precision=3, print_index=True)


class VersionControl(Generator):
    """table with document version control. Prefills first row using orchestrator metadata.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "version_control"
    DEFAULT_HEADING = "Document Version Control"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.version = "1"
        self.author_name = self.orchestrator.metadata["author_name"]
        self.current_date = date.today().strftime("%B %d, %Y")
        if self.author_name == "Triss Merigold":  # easter egg :P shall delete it later... ## why, let's have more characters added here :-)
            self.details = "Casted spell to make documentation appear 🧙‍♀️"
        else:
            self.details = "Automatically generated documentation"

    def calculate(self):

        self.version_control_table = pd.DataFrame({
            "Version": [self.version, "", "", "", ""],
            "Date": [self.current_date, "", "", "", ""],
            "Author": [self.author_name, "", "", "", ""],
            "Change Details": [self.details, "", "", "", ""],
        })

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.table_from_dataframe(self.version_control_table, style=None, print_index=False)


class BasicInformation(Generator):
    """ Adds basic info from metadata dictionary - country, area and segment, into docx.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
    """
    ID = "basic_information"
    DEFAULT_HEADING = "Basic Information"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.country = self.orchestrator.metadata["country"]
        self.area = self.orchestrator.metadata["area"]
        self.segment = self.orchestrator.metadata["segment"]

    def calculate(self):
        """ pass """
        pass

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"Country: {self.country}\n")
        p.add_run(f"Area: {self.area}\n")
        p.add_run(f"Segment: {self.segment}\n")


class TargetDefinition(Generator):
    """Prints target name.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "target_definition"
    DEFAULT_HEADING = "Target Definition"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.target = self.orchestrator.metadata["columns"]["target"]

    def calculate(self):
        pass

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"Name: {self.target}\n")
        p.add_run(f"Business definition: \n")


class ScoringScriptSql(Generator):
    """SQL script to score new data with Logistic regression PSW scorecard on a Oracle DWH.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "scoring_script_sql"
    DEFAULT_HEADING = "Scoring SQL Script"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.target = self.columns["target"]
        self.grouping_path = self.orchestrator.metadata["grouping_path"]
        self.model_path = self.orchestrator.metadata["model_path"]

    def calculate(self):
        gr = grouping_new.interactive.NewGrouping()
        gr.init_underlying_grouping(self.data[self.predictors])
        gr.load(self.grouping_path)
        md = pickle.load(open(self.model_path, 'rb'))
        sc = scorecard.ScoreCard(
            grouping=gr.grouping,
            predictors=md.predictors,
            coefficients=md.coef,
            intercept=md.intercept,
        )

        self.code = sc.to_SQL(
            ntbOut=True
        )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(self.code)


class ScoringScriptPythonGM(Generator):
    """Python script to score new data using the General model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "scoring_script_python_gm"
    DEFAULT_HEADING = "Scoring Python Script"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.model_path = self.orchestrator.metadata["model_path"]

    def calculate(self):
        md = pickle.load(open(self.model_path, 'rb'))
        self.code = md.transformation_code(dataset_name='data')

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(self.code)


class PredictorName(Generator):
    """Outputs predictor name as a heading of a chapter.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        predictor_name (str): name of predictor. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "predictor_name"

    def calculate(self):
        pass

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.predictor_name, level=self.level)


class PredictorStability(Generator):
    """Stability of predictor distribution and bad rate in each bin of the predictor.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
        predictor_name (str): name of predictor. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "predictor_stability"
    DEFAULT_HEADING = "Predictor Stability"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.target = self.columns["target"]
        self.base = self.columns["base"]

        if "grouping_path" in self.orchestrator.metadata:
            self.grouping_path = self.orchestrator.metadata["grouping_path"]
        else:
            self.grouping_path = None
        self.data_train = self.orchestrator.data[self.orchestrator.metadata["samples"]["Train"]]
        self.data = self.orchestrator.data[self.orchestrator.metadata["samples"]["All"]]

        self.output_folder_p = os.path.join(self.output_folder, "predictor_stability")

        if self.use_weight:
            self.output_folder_weight = os.path.join(self.output_folder, "predictor_stability_weighted")

    def calculate(self):
        if not os.path.exists(self.output_folder_p):
            os.makedirs(self.output_folder_p)
        self.plot_path = self.output_folder_p + '/' + self.predictor_name + '_WOE.PNG'

        if self.use_weight:
            if not os.path.exists(self.output_folder_weight):
                os.makedirs(self.output_folder_weight)
            self.plot_path_weight = self.output_folder_weight + '/' + self.predictor_name + '_WOE.PNG'

        gr = grouping_new.interactive.NewGrouping()
        gr.init_underlying_grouping(self.data[[self.predictor_name]])
        if self.grouping_path:
            gr.load(self.grouping_path)
        else:
            bin_dict = features.fake_binning(self.data_train[[self.predictor_name]], bin_count=5)
            self.grouping_path = 'documentator_fake_binning.json'
            with open(self.grouping_path, 'w', encoding='utf-8') as file:
                json.dump(bin_dict, file, ensure_ascii=False, cls=grouping.NumpyJSONEncoder, indent=2)
        gr.load(self.grouping_path)

        transformed_data = gr.transform(self.data[[self.predictor_name]])
        self.data = self.data.join(transformed_data, lsuffix='_ORIG')

        self.orchestrator.doctools.GroupedEvaluation(
            data=self.data,
            predictor=self.predictor_name + '_WOE',
            sample="All",
            target=self.target,
            grouping=gr.grouping,
            use_weight=False,
            output_folder=self.output_folder_p,
            display_table=False,
            show_plot=False,
        )

        if self.use_weight:
            self.orchestrator.doctools.GroupedEvaluation(
                data=self.data,
                predictor=self.predictor_name + '_WOE',
                sample="All",
                target=self.target,
                grouping=gr.grouping,
                use_weight=True,
                output_folder=self.output_folder_weight,
                display_table=False,
                show_plot=False,
            )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_picture(self.plot_path, width=docx.shared.Inches(self.PIC_WIDTH_L))
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Weighted:")
            self.orchestrator.document.add_picture(self.plot_path_weight, width=docx.shared.Inches(self.PIC_WIDTH_L))



class PredictorStabilityShortTarget(PredictorStability):
    """Stability of predictor distribution and bad rate in each bin of the predictor. Calculated using short target from the metadata.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
        predictor_name (str): name of predictor. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "predictor_stability_short_target"
    DEFAULT_HEADING = "Predictor Stability - Short Target"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.target = self.columns["short_target"]
        self.base = self.columns["short_base"]

        self.output_folder_p = os.path.join(self.output_folder, "predictor_stability_short")


class ScorecardHistogram(Generator):
    """
    Returns a scorecard histograms, pd and linear.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
        predictor_name (str): name of predictor. Defaults to None.

    """
    ID = "scorecard_histogram"
    DEFAULT_HEADING = "Scorecard Histogram"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.score = self.columns["score"]
        self.target = self.columns["target"]
        self.base = self.columns["base"]

        self.sample = "Observable"
        self.plot_path_linear = "sc_hist_linear.png"
        self.plot_path_pd = "sc_hist_pd.png"
        self.plot_path_linear_w = "sc_hist_linear_w.png"
        self.plot_path_pd_w = "sc_hist_pd_w.png"

    def calculate(self):
        """ prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.ScoreHistogram`"""
        self.orchestrator.doctools.ScoreHistogram(data=self.data,
                                                  sample=self.sample,
                                                  target=self.target,
                                                  score=self.score,
                                                  use_weight=None,
                                                  use_logit=True,
                                                  output_folder=self.output_folder,
                                                  filename=self.plot_path_linear,
                                                  show_plot=False
                                                  )
        self.orchestrator.doctools.ScoreHistogram(data=self.data,
                                                  sample=self.sample,
                                                  target=self.target,
                                                  score=self.score,
                                                  use_weight=None,
                                                  use_logit=False,
                                                  min_score=0,
                                                  max_score=1,
                                                  output_folder=self.output_folder,
                                                  filename=self.plot_path_pd,
                                                  show_plot=False
                                                  )

        if self.use_weight:
            self.orchestrator.doctools.ScoreHistogram(data=self.data,
                                                      sample=self.sample,
                                                      target=self.target,
                                                      score=self.score,
                                                      use_weight=self.use_weight,
                                                      use_logit=True,
                                                      output_folder=self.output_folder,
                                                      filename=self.plot_path_linear_w,
                                                      show_plot=False
                                                      )
            self.orchestrator.doctools.ScoreHistogram(data=self.data,
                                                      sample=self.sample,
                                                      target=self.target,
                                                      score=self.score,
                                                      use_weight=self.use_weight,
                                                      use_logit=False,
                                                      min_score=0,
                                                      max_score=1,
                                                      output_folder=self.output_folder,
                                                      filename=self.plot_path_pd_w,
                                                      show_plot=False
                                                      )

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"Score:\t\t{self.score}\n")
        p.add_run(f"Sample:\t{self.sample}\n")
        p.add_run(f"Target:\t\t{self.target}\n")
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_linear), width=docx.shared.Inches(self.PIC_WIDTH_M),)
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_pd), width=docx.shared.Inches(self.PIC_WIDTH_M),)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Weighted Scorecard Histograms")
            self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_linear_w), width=docx.shared.Inches(self.PIC_WIDTH_M),)
            self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_pd_w), width=docx.shared.Inches(self.PIC_WIDTH_M),)


class CovariateList(Generator):
    """
    Table with all considered covariates and their performance.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
    """
    ID = "covariate_list"
    DEFAULT_HEADING = "List of Covariates"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.masks = ["Train", "Validation", "Test", "Out of Time", "Historical Out of Time"]
        self.target = self.columns["target"]
        self.data_train = self.orchestrator.data[self.orchestrator.metadata["samples"]["Train"]]
        
        if "covariates" in self.orchestrator.metadata.keys():
            self.covariates = self.orchestrator.metadata["covariates"]
        else:
            self.covariates = self.orchestrator.metadata["predictors"]

        if "grouping_path" in self.orchestrator.metadata.keys():
            self.grouping_path = self.orchestrator.metadata["grouping_path"]
        else:
            self.grouping_path = None
        
        self.covariates_transformed = [s + "_WOE" for s in self.covariates]

        self.plot_path = "covariates.csv"
        self.plot_path_w = "covariates_w.csv"

    def calculate(self):
        """ prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.PredictorPowerAnalysis`"""
        gr = grouping_new.interactive.NewGrouping()
        gr.init_underlying_grouping(self.data[self.covariates])
        
        if not self.grouping_path:
            bin_dict = features.fake_binning(self.data_train[self.covariates], bin_count=10)
            self.grouping_path = 'documentator_fake_binning.json'
            with open(self.grouping_path, 'w', encoding='utf-8') as file:
                json.dump(bin_dict, file, ensure_ascii=False, cls=grouping.NumpyJSONEncoder, indent=2)

        gr.load(self.grouping_path)
        transformed_data = gr.transform(self.data[self.covariates])
        self.data = self.data.join(transformed_data, lsuffix='_ORIG')

        self.covariate_list = self.orchestrator.doctools.PredictorPowerAnalysis(data=self.data,
                                                                                sample=self.sample,
                                                                                predictors=self.covariates_transformed,
                                                                                target=self.target,
                                                                                sort_by=None,
                                                                                use_weight=False,
                                                                                masks=self.masks,
                                                                                output_folder=self.output_folder,
                                                                                filename=self.plot_path)
        self.covariate_list = self.covariate_list.droplevel([0], axis=1)
        if self.use_weight:
            self.covariate_list_w = self.orchestrator.doctools.PredictorPowerAnalysis(data=self.data,
                                                                                      sample=self.sample,
                                                                                      predictors=self.covariates_transformed,
                                                                                      target=self.target,
                                                                                      sort_by=None,
                                                                                      use_weight=False,
                                                                                      masks=self.masks,
                                                                                      output_folder=self.output_folder,
                                                                                      filename=self.plot_path_w)
            self.covariate_list_w = self.covariate_list_w.droplevel([0], axis=1)

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.table_from_dataframe(self.covariate_list, style=None, float_precision=3)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted Covariates")
            self.table_from_dataframe(self.covariate_list_w, style=None, float_precision=3)


class CovariateListGM(Generator):
    """Table with all considered covariates and their performance. This version is used for General model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "covariate_list_gm"
    DEFAULT_HEADING = "List of Covariates"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.target = self.columns["target"]
        
        if "covariates" in self.orchestrator.metadata.keys():
            self.covariates = self.orchestrator.metadata["covariates"]
        else:
            self.covariates = self.orchestrator.metadata["predictors"]

        if self.use_weight:
            self.weight = self.columns["weight"]

        self.plot_path = "covariates.csv"
        self.plot_path_w = "covariates_w.csv"

    def calculate(self):
        
        self.uni_ginis = []

        if self.use_weight:
            self.uni_ginis_w = []

        for covariate in self.covariates:

            covariate_gini = np.abs(metrics.gini(
                self.data[(self.orchestrator.metadata["samples"][self.sample]) & pd.notnull(self.data[covariate])][self.target],
                self.data[(self.orchestrator.metadata["samples"][self.sample]) & pd.notnull(self.data[covariate])][covariate],
            ))

            self.uni_ginis.append({'Covariate': covariate, 'Gini': covariate_gini})

            if self.use_weight:

                covariate_gini_w = np.abs(metrics.gini(
                    self.data[(self.orchestrator.metadata["samples"][self.sample]) & pd.notnull(self.data[covariate])][self.target],
                    self.data[(self.orchestrator.metadata["samples"][self.sample]) & pd.notnull(self.data[covariate])][covariate],
                    self.data[(self.orchestrator.metadata["samples"][self.sample]) & pd.notnull(self.data[covariate])][self.weight],
                ))

                self.uni_ginis_w.append({'Covariate': covariate, 'Gini': covariate_gini_w})

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.table_from_dataframe(pd.DataFrame(self.uni_ginis), style=None, float_precision=3, print_index=False)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted Covariates")
            self.table_from_dataframe(pd.DataFrame(self.uni_ginis_w), style=None, float_precision=3, print_index=False)


class CorrelationMatrix(Generator):
    """
    Correlation matrix for given list of covariates.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
    """
    ID = "correlation_matrix"
    DEFAULT_HEADING = "Correlation Matrix"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.plot_path = "correlation_matrix.png"
        self.plot_path_w = "correlation_matrix_w.png"
        self.data_train = self.orchestrator.data[self.orchestrator.metadata["samples"]["Train"]]

        if "grouping_path" in self.orchestrator.metadata.keys():
            self.grouping_path = self.orchestrator.metadata["grouping_path"]
        else:
            self.grouping_path = None
        
        self.predictors_woe = [s + "_WOE" for s in self.predictors]

    def calculate(self):
        """ prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.Correlations`"""
        gr = grouping_new.interactive.NewGrouping()
        gr.init_underlying_grouping(self.data[self.predictors])
        
        if not self.grouping_path:
            bin_dict = features.fake_binning(self.data_train[self.predictors], bin_count=10)
            self.grouping_path = 'documentator_fake_binning.json'
            with open(self.grouping_path, 'w', encoding='utf-8') as file:
                json.dump(bin_dict, file, ensure_ascii=False, cls=grouping.NumpyJSONEncoder, indent=2)

        gr.load(self.grouping_path)
        transformed_data = gr.transform(self.data[self.predictors])
        self.data = self.data.join(transformed_data, lsuffix='_ORIG')

        self.orchestrator.doctools.Correlations(data=self.data,
                                                predictors=self.predictors_woe,
                                                sample=self.sample,
                                                use_weight=False,
                                                output_folder=self.output_folder,
                                                filename=self.plot_path,
                                                show_plot=False
                                                )

        if self.use_weight:
            self.orchestrator.doctools.Correlations(data=self.data,
                                                    predictors=self.predictors_woe,
                                                    sample=self.sample,
                                                    use_weight=self.use_weight,
                                                    output_folder=self.output_folder,
                                                    filename=self.plot_path_w,
                                                    show_plot=False
                                                    )

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path), width=docx.shared.Inches(self.PIC_WIDTH_L))
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted Correlation")
            self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_w), width=docx.shared.Inches(self.PIC_WIDTH_L))


class CorrelationMatrixGM(Generator):
    """Correlation matrix. This version is used for General model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "correlation_matrix_gm"
    DEFAULT_HEADING = "Correlation Matrix"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.plot_path = "correlation_matrix.png"
        self.plot_path_w = "correlation_matrix_w.png"
        self.model_path = self.orchestrator.metadata['model_path']

    def calculate(self):
        model = pickle.load(open(self.model_path, 'rb'))
        transformed_data = model.impute(self.data)
        self.data = self.data.join(transformed_data, lsuffix='_ORIG')

        self.orchestrator.doctools.Correlations(data=self.data,
                                                predictors=self.predictors,
                                                sample=self.sample,
                                                use_weight=False,
                                                output_folder=self.output_folder,
                                                filename=self.plot_path,
                                                show_plot=False
                                                )

        if self.use_weight:
            self.orchestrator.doctools.Correlations(data=self.data,
                                                    predictors=self.predictors,
                                                    sample=self.sample,
                                                    use_weight=self.use_weight,
                                                    output_folder=self.output_folder,
                                                    filename=self.plot_path_w,
                                                    show_plot=False
                                                    )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path), width=docx.shared.Inches(self.PIC_WIDTH_L))
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted Correlation")
            self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_w), width=docx.shared.Inches(self.PIC_WIDTH_L))


class PerformanceTable(Generator):
    """
    Table with Gini, Lift and Kolmogorov-Smirnov test for given set of masks (train, validation, test, ...)

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.
    """
    ID = "performance_table"
    DEFAULT_HEADING = "Model Performance"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.mask = ["Train", "Validation", "Test", "Out of Time", "Historical Out of Time"]
        self.filename = "performance.csv"
        self.filename_w = "performance_w.csv"
        self.score = self.columns["score"]
        self.target = self.columns["target"]
        self.lift_perc = 10

    def calculate(self):
        """ prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.ModelPerformanceGiniLiftKS`"""
        self.perf_table = self.orchestrator.doctools.ModelPerformanceGiniLiftKS(data=self.data,
                                                                                sample=self.sample,
                                                                                target=self.target,
                                                                                scores=self.score,
                                                                                lift_perc=self.lift_perc,
                                                                                use_weight=False,
                                                                                masks=self.mask,
                                                                                output_folder=self.output_folder,
                                                                                filename=self.filename
                                                                                )

        if self.use_weight:
            self.perf_table_w = self.orchestrator.doctools.ModelPerformanceGiniLiftKS(data=self.data,
                                                                                      sample=self.sample,
                                                                                      target=self.target,
                                                                                      scores=self.score,
                                                                                      lift_perc=self.lift_perc,
                                                                                      use_weight=self.use_weight,
                                                                                      masks=self.mask,
                                                                                      output_folder=self.output_folder,
                                                                                      filename=self.filename_w
                                                                                      )

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.table_from_dataframe(self.perf_table, style=None, float_precision=3)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted Performance")
            self.table_from_dataframe(self.perf_table_w, style=None, float_precision=3)


class PerformanceTableOld(PerformanceTable):
    """
    Table with Gini, Lift and Kolmogorov-Smirnov test for given set of masks (train, validation, test, ...) and multiple
    scores (usually current and old score).

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.

    Methods:
        calculate(): prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.ModelPerformanceGiniLiftKS`
        export_to_document(): export the output to given Word docx
    """
    ID = "performance_table_oldscore"
    DEFAULT_HEADING = "Model Performance Comparison with Old Score"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.mask = ["Train", "Validation", "Test", "Out of Time", "Historical Out of Time"]
        self.filename = "comparison_performance.csv"
        self.filename_w = "comparison_performance_w.csv"
        self.score = [self.columns["old_score"], self.columns['score']]
        self.target = self.columns["target"]
        self.lift_perc = 10


class PerformanceBootstrap(Generator):
    """
    Bootstrapped model's Gini. Mean, Standard Deviation, Confidence intervals.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.
    """
    ID = "performance_bootstrap"
    DEFAULT_HEADING = "Model Performance Bootstrapped"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.mask = ["Train", "Validation", "Test", "Out of Time", "Historical Out of Time"]
        self.filename = "gini_bootstrap.csv"
        self.filename_w = "gini_bootstrap_w.csv"
        self.score = self.columns["score"]
        self.target = self.columns["target"]
        # self.lift_perc = 10
        self.random_seed = 1234

    def calculate(self):
        """ prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.ScoreGiniBootstrap`"""
        self.perfbootstrap = self.orchestrator.doctools.ScoreGiniBootstrap(data=self.data,
                                                                           sample="Observable",
                                                                           scores=self.score,
                                                                           target=self.target,
                                                                           masks=self.mask,
                                                                           use_weight=False,
                                                                           n_iter=100,
                                                                           ci_range=5,
                                                                           random_seed=self.random_seed,
                                                                           # col_score_ref=col_oldscore,
                                                                           output_folder=self.output_folder,
                                                                           filename=self.filename
                                                                           )

        if self.use_weight:
            self.perfbootstrap_w = self.orchestrator.doctools.ScoreGiniBootstrap(data=self.data,
                                                                                 sample="Observable",
                                                                                 scores=self.score,
                                                                                 target=self.target,
                                                                                 masks=self.mask,
                                                                                 use_weight=self.use_weight,
                                                                                 n_iter=100,
                                                                                 ci_range=5,
                                                                                 random_seed=self.random_seed,
                                                                                 # col_score_ref=col_oldscore,
                                                                                 output_folder=self.output_folder,
                                                                                 filename=self.filename_w
                                                                                 )

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.table_from_dataframe(self.perfbootstrap, style=None, float_precision=3, print_index=False)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted Gini Bootstrapped")
            self.table_from_dataframe(self.perfbootstrap_w, style=None, float_precision=3, print_index=False)


class PerformanceBootstrapOld(PerformanceBootstrap):
    """Bootstrapped model's Gini. Mean, Standard Variance, Confidence intervals. Comparison of new and old score.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.

    Methods:
        calculate(): prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.ScoreGiniBootstrap`
        export_to_document(): export the output to given Word docx
    """
    ID = "performance_bootstrap_oldscore"
    DEFAULT_HEADING = "Model Performance Bootstrapped"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.mask = ["Train", "Validation", "Test", "Out of Time", "Historical Out of Time"]
        self.filename = "gini_bootstrap_old.csv"
        self.filename_w = "gini_bootstrap_old_w.csv"
        self.score = [self.columns["score"], self.columns["old_score"]]
        self.target = self.columns["target"]
        # self.lift_perc = 10


class ROCCurves(Generator):
    """
    ROC curve - use for plotting one score and multiple masks.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
        predictor_name (str, optional): name of predictor. Defaults to None.

    """

    ID = "roc_curves"
    DEFAULT_HEADING = "ROC Curves"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.mask = ["Train", "Validation", "Test", "Out of Time", "Historical Out of Time"]
        self.plot_path = "ROCCurve.png"
        self.plot_path_w = "ROCCurve_w.png"
        self.score = self.columns["score"]  
        self.target = self.columns["target"]

    def calculate(self):
        """prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.ROCCurve`"""
        self.orchestrator.doctools.ROCCurve(data=self.data,
                                            sample=self.sample,
                                            scores=self.score,
                                            target=self.target,
                                            masks=self.mask,
                                            use_weight=False,
                                            output_folder=self.output_folder,
                                            filename=self.plot_path,
                                            show_plot=False
                                            )
        if self.use_weight:
            self.orchestrator.doctools.ROCCurve(data=self.data,
                                                sample=self.sample,
                                                scores=self.score,
                                                target=self.target,
                                                masks=self.mask,
                                                use_weight=self.use_weight,
                                                output_folder=self.output_folder,
                                                filename=self.plot_path_w,
                                                show_plot=False
                                                )

    def export_to_document(self):
        """ exports the output to given Word docx """
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path), width=docx.shared.Inches(self.PIC_WIDTH_S))
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted ROC Curve")
            self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_w), width=docx.shared.Inches(self.PIC_WIDTH_S))


class ROCCurvesOld(ROCCurves):
    """
    ROC curve - use for plotting multiple scores on the same sample (e.g. New and Old Score).

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx

    """
    ID = "roc_curves_oldscore"
    DEFAULT_HEADING = "ROC Curves Comparison with Old Score"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Old comparison"
        self.mask = None
        self.plot_path = "ROCCurve_comparison.png"
        self.plot_path_w = "ROCCurve_comparison_w.png"
        self.score = [self.columns["score"], self.columns["old_score"]]


class LiftCurves(Generator):
    """
    Lift

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
        predictor_name (str, optional): name of predictor. Defaults to None.

    """
    ID = "lift_curves"
    DEFAULT_HEADING = "Lift Curves"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.mask = ["Train", "Validation", "Test", "Out of Time", "Historical Out of Time"]
        self.plot_path = "LiftCurve.png"
        self.plot_path_w = "LiftCurve_w.png"
        self.score = self.columns["score"]
        self.target = self.columns["target"]
        self.lift_perc = 15

    def calculate(self):
        """
        prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.LiftCurve`
        """
        self.orchestrator.doctools.LiftCurve(data=self.data,
                                             sample=self.sample,
                                             scores=self.score,
                                             target=self.target,
                                             masks=self.mask,
                                             use_weight=False,
                                             lift_perc=self.lift_perc,
                                             output_folder=self.output_folder,
                                             filename=self.plot_path,
                                             show_plot=False
                                             )
        if self.use_weight:
            self.orchestrator.doctools.LiftCurve(data=self.data,
                                                 sample=self.sample,
                                                 scores=self.score,
                                                 target=self.target,
                                                 masks=self.mask,
                                                 use_weight=self.use_weight,
                                                 lift_perc=self.lift_perc,
                                                 output_folder=self.output_folder,
                                                 filename=self.plot_path_w,
                                                 show_plot=False
                                                 )

    def export_to_document(self):
        """ exports the output to given Word docx """
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path), width=docx.shared.Inches(self.PIC_WIDTH_S))
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted Lift Curve")
            self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_w), width=docx.shared.Inches(self.PIC_WIDTH_S))


class LiftCurvesOld(LiftCurves):
    """
        Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "lift_curves_oldscore"
    DEFAULT_HEADING = "Lift Curves Comparison with Old Score"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Old comparison"
        self.mask = None
        self.plot_path = "LiftCurve_comparison.png"
        self.plot_path_w = "LiftCurve_comparison_w.png"
        self.score = [self.columns["score"], self.columns["old_score"]]


class GiniLiftInTime(Generator):
    """
    Creates plots of Gini AND Lift (two plots) in time for one score, one target and multiple samples (train, test, valid...)
    The Lift is set to 15 %.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.
    """
    ID = "gini_lift_intime"
    DEFAULT_HEADING = "Gini in Time"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.mask = ["Train", "Validation", "Test", "Out of Time", "Historical Out of Time"]
        self.plot_path_g = "gini_in_time.png"
        self.plot_path_gw = "gini_in_time_w.png"
        self.plot_path_l = "lift_in_time.png"
        self.plot_path_lw = "lift_in_time_w.png"
        self.score = self.columns["score"]
        self.target = self.columns["target"]
        self.lift_perc = 15

        self.heading_lift = "Lift in Time"

    def calculate(self):
        """
        prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.GiniLiftInTimeScore`
        """
        self.orchestrator.doctools.GiniLiftInTimeScore(data=self.data,
                                                       sample=self.sample,
                                                       scores=self.score,
                                                       masks=self.mask,
                                                       target=self.target,
                                                       use_weight=False,
                                                       lift_perc=self.lift_perc,
                                                       get_gini=True,
                                                       get_lift=True,
                                                       output_folder=self.output_folder,
                                                       filename_gini=self.plot_path_g,
                                                       filename_lift=self.plot_path_l,
                                                       show_plot=False
                                                       )
        if self.use_weight:
            self.orchestrator.doctools.GiniLiftInTimeScore(data=self.data,
                                                           sample=self.sample,
                                                           scores=self.score,
                                                           masks=self.mask,
                                                           target=self.target,
                                                           use_weight=self.use_weight,
                                                           lift_perc=self.lift_perc,
                                                           get_gini=True,
                                                           get_lift=True,
                                                           output_folder=self.output_folder,
                                                           filename_gini=self.plot_path_gw,
                                                           filename_lift=self.plot_path_lw,
                                                           show_plot=False
                                                           )

    def export_to_document(self):
        """ exports the output to given Word docx """
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_g), width=docx.shared.Inches(self.PIC_WIDTH_M))

        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted Gini in Time")
            self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_gw), width=docx.shared.Inches(self.PIC_WIDTH_M))

        self.orchestrator.document.add_heading(self.heading_lift, level=self.level)
        self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_l), width=docx.shared.Inches(self.PIC_WIDTH_M))

        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run("\n")
            p.add_run(f"Weighted Lift in Time")
            self.orchestrator.document.add_picture(os.path.join(self.output_folder, self.plot_path_lw), width=docx.shared.Inches(self.PIC_WIDTH_M))


class GiniLiftInTimeShortTgt(GiniLiftInTime):
    """
    Creates plots of Gini AND Lift (two plots) in time for one score, one sample and multiple targets (e.g. target and short target)
    The Lift is set to 15 %.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.

    Methods:
        calculate(): prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.GiniLiftInTimeScore`
        export_to_document(): export the output to given Word docx"""
    ID = "gini_lift_intime_short"
    DEFAULT_HEADING = "Gini in Time for Short Target"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Observable"
        self.mask = None
        self.plot_path_g = "gini_in_time_short.png"
        self.plot_path_gw = "gini_in_time_short_w.png"
        self.plot_path_l = "lift_in_time_short.png"
        self.plot_path_lw = "lift_in_time_short_w.png"
        self.target = [self.columns["target"], self.columns["short_target"]]
        self.lift_perc = 15
        self.heading_lift = "Lift in Time for Short Target"


class GiniLiftInTimeOld(GiniLiftInTime):
    """
    Creates plots of Gini AND Lift (two plots) in time for one target, one sample and multiple scores (e.g. new and old)
    The Lift is set to 15 %.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.

    Methods:
        calculate(): prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.GiniLiftInTimeScore`
        export_to_document(): export the output to given Word docx
    """
    ID = "gini_lift_intime_oldscore"
    DEFAULT_HEADING = "Gini in Time Comparison"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Old comparison"
        self.mask = None
        self.plot_path_g = "gini_in_time_old.png"
        self.plot_path_gw = "gini_in_time_old_w.png"
        self.plot_path_l = "lift_in_time_old.png"
        self.plot_path_lw = "lift_in_time_old_w.png"
        self.score = [self.columns["score"], self.columns["old_score"]]
        self.target = self.columns["target"]
        self.lift_perc = 15
        self.heading_lift = "Lift in Time Comparison"


class GiniLiftInTimeOldShort(GiniLiftInTime):
    """
    Creates plots of Gini AND Lift (two plots) in time for multiple scores AND targets, on one sample.
    The Lift is set to 15 %.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. Defaults to None.

    Methods:
        calculate(): prepares the output, calling :py:meth:`~scoring.doctools.calculators.ProjectParameters.GiniLiftInTimeScore`
        export_to_document(): export the output to given Word docx
    """
    ID = "gini_lift_intime_short_target_oldscore"
    DEFAULT_HEADING = "Gini in Time Comparison with Short Target"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.sample = "Old comparison"
        self.mask = None
        self.plot_path_g = "gini_in_time_oldshort.png"
        self.plot_path_gw = "gini_in_time_oldshort_w.png"
        self.plot_path_l = "lift_in_time_oldshort.png"
        self.plot_path_lw = "lift_in_time_oldshort_w.png"
        self.score = [self.columns["score"], self.columns["old_score"]]
        self.target = [self.columns["target"], self.columns["short_target"]]
        self.lift_perc = 15
        self.heading_lift = "Lift in Time Comparison with Short Target"


class PredictorGrouping(Generator):
    """Grouping charts for a predictor. Shows fine classing and coarse classing.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
        predictor_name (str): name of predictor. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """

    ID = "predictor_grouping"
    DEFAULT_HEADING = "Predictor Grouping"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.target = self.columns["target"]
        self.model_path = self.orchestrator.metadata["model_path"]

        self.target = self.columns["target"]
        self.time_column = self.columns["time"]

        self.grouping_path = self.orchestrator.metadata["grouping_path"]
        self.data = self.orchestrator.data[self.orchestrator.metadata["samples"]["Train"]]

        if self.use_weight:
            self.weight = self.columns["weight"]
            self.weight_series = self.data[self.weight]
        else:
            self.weight = None
            self.weight_series = None

        self.output_folder_p = os.path.join(self.output_folder, "predictor_grouping")

    def calculate(self):

        if not os.path.exists(self.output_folder_p):
            os.makedirs(self.output_folder_p)
        self.plot_path = os.path.join(self.output_folder_p, self.predictor_name + '.png')

        gr = grouping_new.interactive.NewGrouping()
        gr.init_underlying_grouping(self.data[[self.predictor_name]])
        gr.load(self.grouping_path)

        gr.export_pictures(
            data=self.data[[self.predictor_name]],
            target=self.data[self.target],
            time_column=self.data[self.time_column],
            weight=self.weight_series,
            export_path=self.output_folder_p,
            use_tqdm=False,
        )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Grouping was performed using weighted data.\n")
        self.orchestrator.document.add_picture(self.plot_path, width=docx.shared.Inches(self.PIC_WIDTH_L))


class PredictorPDP(Generator):
    """Partial dependency plot for a predictor. Suitable for LGBM workflow model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
        predictor_name (str): name of predictor. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """

    ID = "predictor_pdp"
    DEFAULT_HEADING = "Predictor Partial Dependency Plot"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        with open(self.orchestrator.metadata['model_path'], 'rb') as openfile:
            self.lgb_wrapper = pickle.load(openfile)
        model = self.lgb_wrapper.models[0]
        self.orchestrator.doctools.model = (model,'LGBM',model.feature_name())
        self.target = self.columns["target"]

        self.output_folder_p = os.path.join(self.output_folder, "pdp")

        if self.use_weight:
            self.output_folder_weight = os.path.join(self.output_folder, "pdp_weighted")

    def calculate(self):
        if not os.path.exists(self.output_folder_p):
            os.makedirs(self.output_folder_p)
        self.plot_path = os.path.join(self.output_folder_p, 'pdp_' + self.predictor_name + '.PNG')

        if self.use_weight:
            if not os.path.exists(self.output_folder_weight):
                os.makedirs(self.output_folder_weight)
            self.plot_path_weight = self.output_folder_weight + '/pdp_' + self.predictor_name + '.PNG'

        self.orchestrator.doctools.PartialDependencePlot(
            data=self.data,
            sample='Test',
            target=self.target,
            predictor=self.predictor_name,
            output_folder=self.output_folder_p,
            show_plot=False,
            show_table=False,
            use_weight=False,
        )

        if self.use_weight:
            self.orchestrator.doctools.PartialDependencePlot(
                data=self.data,
                sample='Test',
                target=self.target,
                predictor=self.predictor_name,
                output_folder=self.output_folder_weight,
                show_plot=False,
                show_table=False,
                use_weight=True,
            )

    def export_to_document(self):
        """exports the output to given Word docx"""
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_picture(self.plot_path, width=docx.shared.Inches(self.PIC_WIDTH_M))
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Weighted:")
            self.orchestrator.document.add_picture(self.plot_path_weight, width=docx.shared.Inches(self.PIC_WIDTH_M))


class PredictorDistributionCalibration(Generator):
    """Distribution and calibration chart for a predictor.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
        predictor_name (str): name of predictor. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """

    ID = "predictor_distribution_calibration"
    DEFAULT_HEADING = "Predictor Distribution and Calibration"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)
        self.target = self.columns["target"]

        self.output_folder_p = os.path.join(self.output_folder, "predictor_distribution")

        if self.use_weight:
            self.weight = self.columns["weight"]

        if "predictors_pd" in self.orchestrator.metadata.keys():
            self.predictors_pd = self.orchestrator.metadata["predictors_pd"]
        else:
            self.predictors_pd = []

        if "predictors_rescale" in self.orchestrator.metadata.keys():
            self.predictors_rescale = self.orchestrator.metadata["predictors_rescale"]
        else:
            self.predictors_rescale = []

    def calculate(self):
        if not os.path.exists(self.output_folder_p):
            os.makedirs(self.output_folder_p)
        self.plot_path = self.output_folder_p + '/' + self.predictor_name + '.png'
        self.plot_path_weight = self.output_folder_p + '/' + self.predictor_name + '_weighted.png'

        if (self.predictor_name[-3:] == 'WOE') and (self.use_weight):
            self.shift, self.scale = np.log(
                (
                    self.data[self.orchestrator.metadata["samples"]["Observable"]][self.target]*self.data[self.orchestrator.metadata["samples"]["Observable"]][self.weight]
                ).sum() / (
                    (1-self.data[self.orchestrator.metadata["samples"]["Observable"]][self.target])*self.data[self.orchestrator.metadata["samples"]["Observable"]][self.weight]
                ).sum()
            ), -1
        elif (self.predictor_name[-3:] == 'WOE'):
            self.shift, self.scale = np.log(
                (
                    self.data[self.orchestrator.metadata["samples"]["Observable"]][self.target]
                ).sum() / (
                    (1-self.data[self.orchestrator.metadata["samples"]["Observable"]][self.target])
                ).sum()
            ), -1
        elif (self.use_weight) and (self.predictor_name in self.predictors_rescale):
            scaler = LogisticRegression(penalty = 'l2', C = 1000, solver='liblinear')
            scaler.fit(
                self.data[(self.orchestrator.metadata["samples"]["Observable"]) & pd.notnull(self.data[self.predictor_name])][[self.predictor_name]],
                self.data[(self.orchestrator.metadata["samples"]["Observable"]) & pd.notnull(self.data[self.predictor_name])][self.target],
                sample_weight=self.data[(self.orchestrator.metadata["samples"]["Observable"]) & pd.notnull(self.data[self.predictor_name])][self.weight],
            )
            self.shift, self.scale = scaler.intercept_[0], scaler.coef_[0][0]
        elif (self.predictor_name in self.predictors_rescale):
            scaler = LogisticRegression(penalty = 'l2', C = 1000, solver='liblinear')
            scaler.fit(
                self.data[(self.orchestrator.metadata["samples"]["Observable"]) & pd.notnull(self.data[self.predictor_name])][[self.predictor_name]],
                self.data[(self.orchestrator.metadata["samples"]["Observable"]) & pd.notnull(self.data[self.predictor_name])][self.target],
            )
            self.shift, self.scale = scaler.intercept_[0], scaler.coef_[0][0]
        elif (self.use_weight) and metrics.gini(
            self.data[(self.orchestrator.metadata["samples"]["Observable"]) & pd.notnull(self.data[self.predictor_name])][self.target],
            self.data[(self.orchestrator.metadata["samples"]["Observable"]) & pd.notnull(self.data[self.predictor_name])][self.predictor_name],
            sample_weight=self.data[(self.orchestrator.metadata["samples"]["Observable"]) & pd.notnull(self.data[self.predictor_name])][self.weight],
        ) < 0:
            self.shift, self.scale = 0, -1
        elif metrics.gini(
            self.data[(self.orchestrator.metadata["samples"]["Observable"]) & pd.notnull(self.data[self.predictor_name])][self.target],
            self.data[(self.orchestrator.metadata["samples"]["Observable"]) & pd.notnull(self.data[self.predictor_name])][self.predictor_name],
        ) < 0:
            self.shift, self.scale = 0, -1
        else:
            self.shift, self.scale = 0, 1

        ispd = self.predictor_name in self.predictors_pd

        self.orchestrator.doctools.CalibrationDistribution(
            data=self.data,
            sample="Observable",
            target=self.target,
            score=self.predictor_name,
            use_weight=False,
            shift=self.shift,
            scale=self.scale,
            apply_logit=ispd,
            output_folder=self.output_folder_p,
            filename=f'{self.predictor_name}.png',
        )

        if self.use_weight:
            self.orchestrator.doctools.CalibrationDistribution(
                data=self.data,
                sample="Observable",
                target=self.target,
                score=self.predictor_name,
                use_weight=self.use_weight,
                shift=self.shift,
                scale=self.scale,
                apply_logit=ispd,
                output_folder=self.output_folder_p,
                filename=f'{self.predictor_name}_weighted.png',
            )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_picture(self.plot_path, width=docx.shared.Inches(self.PIC_WIDTH_S))
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Weighted:")
            self.orchestrator.document.add_picture(self.plot_path_weight, width=docx.shared.Inches(self.PIC_WIDTH_S))


class PredictorHitRateInTime(Generator):
    """Hit rate chart of a predictor. Shows how hit rate develops in time.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.
        predictor_name (str): name of predictor. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """

    ID = "predictor_hit_rate_in_time"
    DEFAULT_HEADING = "Predictor Hit Rate in Time"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.output_folder_p = os.path.join(self.output_folder, "predictor_hit_rate")

        if self.use_weight:
            self.weight = self.columns["weight"]
            self.output_folder_weight = os.path.join(self.output_folder, "predictor_hit_rate_weighted")

    def calculate(self):
        if not os.path.exists(self.output_folder_p):
            os.makedirs(self.output_folder_p)
        self.plot_path = self.output_folder_p + '/' + self.predictor_name + '.png'

        if self.use_weight:
            if not os.path.exists(self.output_folder_weight):
                os.makedirs(self.output_folder_weight)
            self.plot_path_weight = self.output_folder_weight + '/' + self.predictor_name + '.png'

        self.orchestrator.doctools.EmptyInTime(
            data=self.data,
            sample="Observable",
            predictors=[self.predictor_name],
            use_weight=False,
            output_folder=self.output_folder_p,
            show_plot=False
        )

        if self.use_weight:
            self.orchestrator.doctools.EmptyInTime(
                data=self.data,
                sample="Observable",
                predictors=[self.predictor_name],
                use_weight=self.use_weight,
                output_folder=self.output_folder_weight,
                show_plot=False
            )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_picture(self.plot_path, width=docx.shared.Inches(self.PIC_WIDTH_S))
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Weighted:")
            self.orchestrator.document.add_picture(self.plot_path_weight, width=docx.shared.Inches(self.PIC_WIDTH_S))


class MonotonicContraintsLGBM(Generator):
    """Lists monotonic constraints of LGBM workflow model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """

    ID = "monotonic_constraints_lgbm"
    DEFAULT_HEADING = "Monotonic Constraints"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        with open(self.orchestrator.metadata['model_path'], 'rb') as openfile:
            self.lgb_wrapper = pickle.load(openfile)
        model = self.lgb_wrapper.models[0]
        self.orchestrator.doctools.model = (model,'LGBM',model.feature_name())

        constraints_tuple = tuple(0 for pred in model.feature_name())
        if hasattr(self.lgb_wrapper, 'params'):
            if 'monotone_constraints' in self.lgb_wrapper.params.keys():
                constraints_tuple = self.lgb_wrapper.params['monotone_constraints']
        self.constraints = dict(zip(model.feature_name(), constraints_tuple))

    def calculate(self):
        pass

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        empty = True
        for predictor, value in self.constraints.items():
            if value == 0:
                pass
            elif value == 1:
                p.add_run(f"{predictor}: increasing\n")
                empty = False
            elif value == -1:
                p.add_run(f"{predictor}: decreasing\n")
                empty = False
        if empty:
                p.add_run(f"No constraints applied.")


class ShapLGBM(Generator):
    """Shap charts for a LGBM workflow model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """

    ID = "shap_chart_lgbm"
    DEFAULT_HEADING = "SHAP chart"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        with open(self.orchestrator.metadata['model_path'], 'rb') as openfile:
            self.lgb_wrapper = pickle.load(openfile)

        self.target = self.columns["target"]
        self.data_train = self.orchestrator.data[self.orchestrator.metadata["samples"]["Train"]]
        self.data_valid = self.orchestrator.data[self.orchestrator.metadata["samples"]["Validation"]]
        self.data_test = self.orchestrator.data[self.orchestrator.metadata["samples"]["Test"]]

        self.cols_num, self.cols_cat = [], []
        for name, col in self.orchestrator.data[self.predictors].iteritems():
            if pd.api.types.is_numeric_dtype(col.dtype):
                self.cols_num.append(name)
            else:
                self.cols_cat.append(name)

        if self.use_weight:
            self.weight_series_train = self.data_train[self.columns["weight"]]
            self.weight_series_valid = self.data_valid[self.columns["weight"]]
        else:
            self.weight_series_train = None
            self.weight_series_valid = None

        self.output_folder_s = os.path.join(self.output_folder, "shap")

    def calculate(self):
        if not os.path.exists(self.output_folder_s):
            os.makedirs(self.output_folder_s)

        self.lgb_wrapper.print_shap_values(
            self.cols_num, 
            self.cols_cat, 
            self.data_train,
            self.data_valid, 
            self.data_train[self.target],
            self.data_valid[self.target],
            set_to_shap=self.data_test,
            w_train=self.weight_series_train,
            w_valid=self.weight_series_valid,
            output_folder=self.output_folder_s,
            )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.orchestrator.document.add_picture(self.output_folder_s+'/shap.png', width=docx.shared.Inches(self.PIC_WIDTH_M))
        self.orchestrator.document.add_picture(self.output_folder_s+'/shap_abs.png', width=docx.shared.Inches(self.PIC_WIDTH_M))


class MarginalContributionRemoveLGBM(Generator):
    """Table with marginal contribution of each predictor of the model.
    Version for LGBM workflow model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "marginal_contribution_remove_lgbm"
    DEFAULT_HEADING = "Marginal Contribution of Predictors"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        with open(self.orchestrator.metadata['model_path'], 'rb') as openfile:
            self.lgb_wrapper = pickle.load(openfile)

        self.target = self.columns["target"]
        self.data_train = self.orchestrator.data[self.orchestrator.metadata["samples"]["Train"]]
        self.data_valid = self.orchestrator.data[self.orchestrator.metadata["samples"]["Validation"]]
        self.data_test = self.orchestrator.data[self.orchestrator.metadata["samples"]["Test"]]

        if self.use_weight:
            self.weight_series_train = self.data_train[self.columns["weight"]]
            self.weight_series_valid = self.data_valid[self.columns["weight"]]
            self.weight_series_test = self.data_test[self.columns["weight"]]
        else:
            self.weight_series_train = None
            self.weight_series_valid = None
            self.weight_series_test = None

    def calculate(self):

        self.mc_table = self.lgb_wrapper.marginal_contribution(
            self.data_train,
            self.data_valid, 
            self.data_train[self.target],
            self.data_valid[self.target],
            self.data_test,
            self.data_test[self.target],
            w_train=self.weight_series_train,
            w_valid=self.weight_series_valid,
            set_to_test_weight=self.weight_series_test,
            )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Scorecard evaluation was performed using weighted data.\n")
        self.table_from_dataframe(self.mc_table, style=None, float_precision=3, print_index=False)
        self.table_from_dataframe(self.mc_table, style=None, float_precision=3, print_index=False)


class TransitionMatrices(Generator):
    """Transition matrices showing how distribution of observations and bad rate relate to deciles of old score and new score.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "transition_matrices"
    DEFAULT_HEADING = "Transition Matrices"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.score = self.columns["score"]
        self.old_score = self.columns["old_score"]
        self.target = self.columns["target"]

        self.output_folder_t = os.path.join(self.output_folder, 'transition_matrices')

        if self.use_weight:
            self.weight = self.columns["weight"]

    def calculate(self):
        if not os.path.exists(self.output_folder_t):
            os.makedirs(self.output_folder_t)

        self.orchestrator.doctools.TransitionMatrix(
            data=self.data,
            score_new=self.score,
            score_old=self.old_score,
            sample='Old comparison',
            target=self.target,
            use_weight=False,
            quantiles_count=10,
            show_plot=False,
            draw_default_matrix=True,
            draw_transition_matrix=True,
            output_folder=self.output_folder_t,
            filename_default='default_rates.png',
            filename_transition='transition.png',
        )

        self.orchestrator.doctools.TransitionMatrix(
            data=self.data,
            score_new=self.score,
            score_old=self.old_score,
            sample='Old comparison with rejected',
            target=self.target,
            use_weight=False,
            quantiles_count=10,
            show_plot=False,
            draw_default_matrix=False,
            draw_transition_matrix=True,
            output_folder=self.output_folder_t,
            filename_transition='transition_with_rejected.png',
        )

        if self.use_weight:

            self.orchestrator.doctools.TransitionMatrix(
                data=self.data,
                score_new=self.score,
                score_old=self.old_score,
                sample='Old comparison',
                target=self.target,
                use_weight=False,
                quantiles_count=10,
                show_plot=False,
                draw_default_matrix=True,
                draw_transition_matrix=True,
                output_folder=self.output_folder_t,
                filename_default='default_rates_weighted.png',
                filename_transition='transition_weighted.png',
            )


            self.orchestrator.doctools.TransitionMatrix(
                data=self.data,
                score_new=self.score,
                score_old=self.old_score,
                sample='Old comparison with rejected',
                target=self.target,
                use_weight=True,
                quantiles_count=10,
                show_plot=False,
                draw_default_matrix=False,
                draw_transition_matrix=True,
                output_folder=self.output_folder_t,
                filename_transition='transition_with_rejected_weighted.png',
            )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"Deafult rate matrix:")
        self.orchestrator.document.add_picture(self.output_folder_t+'/default_rates.png', width=docx.shared.Inches(self.PIC_WIDTH_L))
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"Transition matrix:")
        self.orchestrator.document.add_picture(self.output_folder_t+'/transition.png', width=docx.shared.Inches(self.PIC_WIDTH_L))
        p = self.orchestrator.document.add_paragraph()
        p.add_run(f"Transition matrix with rejected observations:")
        self.orchestrator.document.add_picture(self.output_folder_t+'/transition_with_rejected.png', width=docx.shared.Inches(self.PIC_WIDTH_L))
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Weighted deafult rate matrix:")
            self.orchestrator.document.add_picture(self.output_folder_t+'/default_rates_weighted.png', width=docx.shared.Inches(self.PIC_WIDTH_L))
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Weighted transition matrix:")
            self.orchestrator.document.add_picture(self.output_folder_t+'/transition_weighted.png', width=docx.shared.Inches(self.PIC_WIDTH_L))
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"Weighted transition matrix with rejected observations:")
            self.orchestrator.document.add_picture(self.output_folder_t+'/transition_with_rejected_weighted.png', width=docx.shared.Inches(self.PIC_WIDTH_L))


class ExpectedApprovalRateGM(Generator):
    """Table with expected approval rates within given segments which are specified in orchestrator's metadata (dict keys 'subsegments_ar' and 'reference_ar'). Suitable for General model.

    Args:
        orchestrator (Orchestrator): documentation orchestrator with metadata
        level (int): chapter depth (level in multi-leveled chapter tree)
        heading (str, optional): heading of the chapter. If None, a predefined value is used. Defaults to None.

    Methods:
        calculate(): prepares the output
        export_to_document(): export the output to given Word docx
    """
    ID = "expected_approval_rate_gm"
    DEFAULT_HEADING = "Expected Approval Rate"

    def __init__(self, orchestrator, level, **kwargs):
        super().__init__(orchestrator, level, **kwargs)

        self.score = self.columns["score"]
        
        self.subsegments = self.orchestrator.metadata["subsegments_ar"]
        self.reference_ar = self.orchestrator.metadata["reference_ar"]

        if self.use_weight:
            self.weight = self.columns["weight"]

    def calculate(self):

        self.ar_table = self.orchestrator.doctools.ExpectedApprovalRate(
            data=self.data,
            sample="All",
            score=self.score,
            query_subset=self.subsegments,
            reference_ar=self.reference_ar,
            use_weight=False,
        )

        if self.use_weight:
            self.ar_table_w = self.orchestrator.doctools.ExpectedApprovalRate(
                data=self.data,
                sample="All",
                score=self.score,
                query_subset=self.subsegments,
                reference_ar=self.reference_ar,
                use_weight=self.use_weight,
            )

    def export_to_document(self):
        self.orchestrator.document.add_heading(self.heading, level=self.level)
        self.table_from_dataframe(self.ar_table, style=None, float_precision=3, print_index=False)
        if self.use_weight:
            p = self.orchestrator.document.add_paragraph()
            p.add_run(f"\nWeighted:")
            self.table_from_dataframe(self.ar_table_w, style=None, float_precision=3, print_index=False)

